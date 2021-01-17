import sys
from docx import Document

if len(sys.argv) != 2 and len(sys.argv) != 3:
    print('Arguments Error.')
    sys.exit()
if len(sys.argv) == 2:
    if sys.argv[1] == 'version':
        print('FIB_Maker_Core.\nFill In Blanks Test Paper Maker (Core).\nVersion 1.0.0.\nThis project uses the AGPL-3.0 license.\nSee this license at https://www.gnu.org/licenses/agpl-3.0.txt')
    else:
        print('Arguments Error.')
    sys.exit()

doc = Document(sys.argv[1])
print('Processing...')
count = 0
for para in doc.paragraphs:
    count += 1
    t = para.text  # text
    ul = []  # underline
    while t.find('[') != -1:
        t = para.text  # text
        p = doc.paragraphs[count-1].clear()
        s = t.find('[')  # start
        e = t.find(']')  # end
        c = t[s+1:e]  # choose
        for i in range(len(t[0:s])):
            if i in ul:
                p.add_run(' ').underline = True
            else:
                p.add_run(t[i])
        aw = 0  # a word
        for i in range(len(c)):
            if c[i] == ' ' or i == len(c)-1:
                aw += 1
                for j in range(round(aw*1.5)):
                    p.add_run(' ').underline = True
                    ul.append(s+j)
                if c[i] == ' ':
                    p.add_run(' ')
                aw = 0
            else:
                aw += 1
        p.add_run(t[e+1:len(t)])
        t = t[e:len(t)+1]

doc.save(sys.argv[2])
print('Ok!')
